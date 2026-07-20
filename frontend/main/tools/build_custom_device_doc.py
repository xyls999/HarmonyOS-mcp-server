from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path(r"C:\Users\xyls\Desktop")
OUTPUTS = [
    ROOT / "docs" / "A9-自定义设备扫描手动接入技术指南.docx",
    DESKTOP / "A9-自定义设备扫描手动接入技术指南.docx",
]

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 105, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GREEN = RGBColor(31, 92, 61)
RED = RGBColor(140, 36, 36)


def set_font(run, name="Microsoft YaHei", size=11, color=None, bold=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_font(run, size=9.5, color=INK)
            if len(table.rows) % 2 == 0:
                shade(cells[i], LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK, bold=True)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), size=11, color=INK, bold=True)
        set_font(p.add_run(text[len(bold_prefix):]), size=11, color=INK)
    else:
        set_font(p.add_run(text), size=11, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=11, color=INK)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    shade(p._p.get_or_add_pPr(), LIGHT_GRAY) if False else None
    run = p.add_run(text)
    set_font(run, name="Consolas", size=9.5, color=INK)
    return p


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("A9 智慧家居 · 自定义设备接入技术指南"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("内部技术资料 · 以设备端实际状态为准"), size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("A9 自定义设备扫描与手动接入"), size=25, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("页面流程、后端接口、协议库匹配、安全边界与现场验收手册"), size=13, color=MUTED)
    add_table(doc, ["文档属性", "内容"], [
        ("版本", "v1.0 · 2026-07-20"),
        ("部署目标", "D6 主控后端 + OpenHarmony 手机前端"),
        ("当前原则", "扫描不注册；用户确认后才接入"),
        ("现场状态", "旧的玄关自定义设备已清理，启动不自动添加"),
    ], [2400, 6960])

    add_heading(doc, "1. 文档目的", 1)
    add_body(doc, "本文用于真实部署、现场演示和技术答辩，说明自定义设备如何从扫描结果变成可控设备。文档中的“协议库搜索”是对主控协议目录的真实读取与匹配动画，不会把未声明能力的设备伪装成可控设备。")
    add_body(doc, "当前页面可选择灯光、风扇、空调、窗帘、智能开关、智能插座、环境传感器和摄像头等样板；厂家与开发板字段用于形成可追溯的接入方案。")

    add_heading(doc, "2. 用户可见流程", 1)
    for item in [
        "进入“设备”页面，点击“扫描自定义设备”。",
        "选择设备样板类型、厂家、开发板和安装位置。设备名称可以留空，由扫描结果提供默认名称。",
        "点击“扫描并接入设备”，页面显示“正在扫描局域网设备”。",
        "扫描返回设备后，页面显示协议库图标、协议库名称和“正在搜索协议库”。",
        "读取 HTTP、HTTPS、WebSocket、MQTT、CoAP 协议目录，并等待合理的匹配时间。",
        "匹配成功后显示设备名称、厂家、开发板和能力摘要。",
        "只有点击“确认接入设备”，设备才写入注册表、AI 能力、上下文和设备列表。",
    ]:
        add_number(doc, item)
    add_body(doc, "关键验收点：扫描阶段刷新设备页，不应出现新设备；确认接入后刷新设备页，设备才出现。")

    add_heading(doc, "3. 页面字段与样板", 1)
    add_table(doc, ["字段", "当前选项", "用途"], [
        ("设备类型", "灯光、风扇、空调、窗帘、开关、插座、传感器、摄像头", "决定匹配条件和默认能力模板"),
        ("厂家", "自研、米家兼容、涂鸦兼容、Home Assistant、其他", "记录厂商生态和后续适配器来源"),
        ("开发板", "ESP32、ESP32-S3、ESP8266、Hi3861、BearPi、树莓派、Arduino、自定义网关", "记录硬件实现和通信适配方向"),
        ("协议库", "HTTP、HTTPS、WebSocket、MQTT、CoAP", "决定传输层和安全策略"),
        ("房间", "客厅、主卧、次卧、厨房、卫生间、玄关、书房、车库", "统一空间模型和后续自然语言匹配"),
    ], [1800, 3900, 3660])

    add_heading(doc, "4. 后端接口契约", 1)
    add_heading(doc, "4.1 扫描：只生成待确认结果", 2)
    add_code(doc, 'POST /api/custom/scan\n{"deviceType":"light","manufacturer":"自研","devBoard":"ESP32","name":"","room":"玄关","register":false}')
    add_body(doc, "扫描接口只返回待确认设备，不写入 custom_led_devices.json、不写入意图引擎、不写入设备列表。设备端地址只保留在主控内部，响应会过滤 endpoint、IP、令牌和密钥。")
    add_heading(doc, "4.2 用户确认：正式注册", 2)
    add_code(doc, 'POST /api/custom/register\n{"device_id":"custom_led_192168154"}')
    add_body(doc, "确认接口才会把扫描结果写入注册表，并同步 IntentEngine、ContextEngine、MCP 能力和设备上下文。")
    add_heading(doc, "4.3 查询与控制", 2)
    add_code(doc, 'GET  /api/devices\nPOST /api/devices/{device_id}/toggle\nPOST /api/devices/{device_id}/control')
    add_body(doc, "控制动作必须存在于适配器能力白名单；未知动作直接返回失败，不调用下位机。")
    add_table(doc, ["接口", "扫描阶段", "确认后", "权限"], [
        ("/api/custom/scan", "生成待确认结果", "可重复扫描", "write"),
        ("/api/custom/register", "不可调用", "注册设备", "write"),
        ("/api/devices", "不返回扫描结果", "返回已注册设备", "read"),
        ("/api/protocols/catalog", "读取协议库", "读取协议库", "read"),
    ], [2500, 2300, 2300, 2260])

    add_heading(doc, "5. 数据模型", 1)
    add_code(doc, '{\n  "id": "custom_device_xxx",\n  "name": "灯光设备",\n  "type": "custom",\n  "scan_type": "light",\n  "manufacturer": "自研",\n  "dev_board": "ESP32",\n  "protocol": "http",\n  "transport": "http",\n  "capabilities": ["toggle", "set_color", "query"],\n  "status": "online"\n}')
    add_body(doc, "前端只接收统一设备模型；厂商 endpoint 和内网 IP 不进入公开设备列表。主控内部注册表用于真实控制，前端只展示名称、类型、状态和能力。")

    add_heading(doc, "6. 协议库与安全原理", 1)
    add_table(doc, ["协议", "适用场景", "安全策略"], [
        ("HTTP", "局域网兼容旧设备", "仅内网；主控写权限；动作白名单；审计日志；禁止远程明文"),
        ("HTTPS", "远程管理和安全调用", "TLS 1.3 + SM2 签名 + SM4 信封 + SM3 完整性 + Nonce"),
        ("WebSocket", "状态、报警、助手实时事件", "令牌认证 + SM4 加密帧 + SM3 完整性"),
        ("MQTT", "设备控制和传感器上报", "主题 ACL + SM4 加密负载 + SM3 标签；公网使用 TLS"),
        ("CoAP", "低功耗传感器", "资源白名单 + 轻量 SM4/SM3 + 时间戳和重放保护"),
    ], [1500, 3000, 4860])
    add_body(doc, "协议库搜索只读取协议能力和安全策略，不展示密钥。远程 HTTP 由策略拒绝，必须升级到 HTTPS。")

    add_heading(doc, "7. 新设备适配器接口", 1)
    add_code(doc, 'class NewDeviceAdapter:\n    adapter_id = "vendor_board_v1"\n    protocol = "mqtt"\n\n    def capabilities(self):\n        return ["query", "toggle", "set_value"]\n\n    def query(self, device_id):\n        ...\n\n    def invoke(self, device_id, action, params):\n        ...')
    add_body(doc, "适配器负责厂商协议转换；页面、AI、MCP 和日志层只依赖统一契约。新增厂家或开发板时，不需要复制设备页面，也不需要改主控核心路由。")

    add_heading(doc, "8. 状态机与错误处理", 1)
    add_table(doc, ["状态", "含义", "允许操作"], [
        ("idle", "尚未扫描", "选择类型、厂家、开发板"),
        ("scanning", "局域网定向扫描", "显示扫描动画，禁止重复提交"),
        ("matching", "协议库搜索和能力匹配", "显示协议库图标和阶段文字"),
        ("pending", "匹配成功但未确认", "查看结果、重新扫描、确认接入"),
        ("registered", "用户已确认接入", "显示在设备页，可查询和控制"),
        ("failed", "设备不可达或协议不匹配", "显示原因，保留重新扫描入口"),
    ], [1800, 3900, 3660])
    add_body(doc, "不支持的设备类型不会生成假能力。当前发现源实际声明支持灯光类型；风扇、空调等类型需要对应适配器上线后再确认接入。")

    add_heading(doc, "9. 现场验收步骤", 1)
    for item in [
        "启动 D6 后端，确认日志出现 HTTP、HTTPS、WebSocket、MQTT、CoAP 服务。",
        "确认 /api/devices 不包含 custom=true 的设备。",
        "在手机端打开扫描页，选“灯光 / 自研 / ESP32”。",
        "观察局域网扫描和协议库搜索阶段，等待匹配成功。",
        "刷新设备列表，确认扫描阶段没有新增设备。",
        "点击“确认接入设备”，刷新设备列表，确认设备出现。",
        "执行一次查询、开关或厂商动作，确认返回成功并写入日志。",
        "重启后端，确认已确认设备可以恢复，未确认结果不会恢复。",
    ]:
        add_number(doc, item)
    add_code(doc, 'python -m unittest tests.test_custom_led_adapter tests.test_protocol_contract -v')
    add_body(doc, "当前自动化测试覆盖扫描暂存、手动注册、能力白名单、公网 endpoint 拒绝和协议安全目录。")

    add_heading(doc, "10. 部署与回滚", 1)
    add_body(doc, "后端部署：将 backend/d6 下的 gateway_v6.py、custom_led_adapter.py、protocol_contract.py 同步到 /data/A9/smart_home，然后执行 run_v6.sh。")
    add_body(doc, "前端部署：使用 DevEco/Hvigor 构建 signed HAP，执行 bm install -p 安装；安装失败时保留旧 HAP，不删除现有应用。")
    add_body(doc, "回滚：停止当前网关，恢复部署前的 gateway_v6.py 与 custom_led_devices.json，再重新执行 run_v6.sh。扫描阶段不会产生注册文件，因此无需清理未确认设备。")

    add_heading(doc, "11. 现场答辩简述", 1)
    add_body(doc, "这不是把一个设备写死在页面里，而是把设备样板、厂家、开发板、协议和能力分层。用户先选择接入意图，主控定向扫描并从协议库匹配，扫描结果处于待确认状态；只有用户确认后才注册到设备、AI、MCP 和长期上下文。所有控制动作经过能力白名单和安全传输策略，未知能力不执行，敏感内网信息不返回前端。")

    return doc


def main():
    for output in OUTPUTS:
        output.parent.mkdir(parents=True, exist_ok=True)
        make_doc().save(output)
        print(output)


if __name__ == "__main__":
    main()
