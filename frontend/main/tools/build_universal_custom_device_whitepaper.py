"""Build the generic custom-device integration whitepaper.

This document is intentionally platform-agnostic: it describes how a serious
home/industrial controller can onboard arbitrary devices without pretending
that an unsupported device is connected.
"""
from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tools"))
from build_custom_device_doc import (  # noqa: E402
    BLUE, DARK, INK, MUTED, LIGHT_BLUE, LIGHT_GRAY,
    add_body, add_bullet, add_code, add_heading, add_number, add_table,
    set_font,
)
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

DESKTOP = Path(r"C:\Users\xyls\Desktop")
OUTPUTS = [
    DESKTOP / "A9-自定义设备扫描手动接入技术指南.docx",
    ROOT / "docs" / "A9-自定义设备扫描手动接入技术指南.docx",
]


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 13, 6),
        ("Heading 3", 12, DARK, 9, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("通用自定义设备接入平台 · 技术白皮书"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("架构参考文档 · 以现场设备和协议实际能力为准"), size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("通用自定义设备接入平台"), size=26, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("面向任意厂商、任意硬件、任意主流协议的发现、适配、安全控制与状态同步技术白皮书"), size=13, color=MUTED)
    add_table(doc, ["文档属性", "内容"], [
        ("定位", "通用技术架构与实施规范，不绑定某个品牌、型号或单一项目"),
        ("目标", "让新设备通过适配器、桥接器或标准协议接入统一控制平面"),
        ("核心原则", "真实发现、能力声明、最小权限、可验证状态、失败不伪造成功"),
        ("适用范围", "家庭物联网、楼宇控制、实验室、工厂边缘网关和移动端控制台"),
    ], [2400, 6960])

    add_heading(doc, "1. 先明确：什么叫“任何设备都能接入”", 1)
    add_body(doc, "“任何设备”不是无条件猜测未知协议，而是平台提供统一的设备契约、协议适配器和开发套件。只要设备满足下列任一条件，就可以纳入同一套控制模型：设备公开标准协议；设备提供可调用的厂商 API；设备可以通过串口、GPIO、RS-485、CAN 或无线桥接器访问；厂商愿意提供驱动或设备描述文件。")
    add_body(doc, "平台绝不把未联网、未认证、未声明能力的设备显示成已接入，也不把一个设备的能力臆测成另一个设备的能力。无法验证时状态必须是“未发现 / 离线 / 未认证 / 能力未知”，而不是成功。")
    add_table(doc, ["设备条件", "平台处理", "是否可以控制"], [
        ("标准协议可发现且状态接口可读", "自动识别协议、读取能力、生成候选设备", "通过安全策略和用户确认后可控制"),
        ("厂商 API 可访问但无标准发现", "用户输入或扫描得到地址，使用厂商适配器", "适配器声明哪些动作可用后控制"),
        ("只有串口 / GPIO / RS-485", "接入边缘桥接器，由桥接器转换为统一协议", "桥接器在线且动作有回读时可控制"),
        ("设备未联网或协议未知", "记录扫描失败原因，不创建假设备", "不可控制"),
    ], [2700, 3900, 2760])

    add_heading(doc, "2. 总体架构：发现层、能力层、执行层分离", 1)
    add_body(doc, "建议采用六层架构。每层只承担一种责任，避免把厂商细节写进前端页面或 AI 提示词。")
    add_table(doc, ["层", "职责", "关键输出"], [
        ("设备与网络层", "Wi-Fi、以太网、BLE、Zigbee、Thread、RS-485、CAN、串口等真实连接", "端点、链路质量、在线心跳"),
        ("发现与配网层", "扫描、配对、二维码、用户输入、厂商云或本地目录", "设备候选、认证材料、发现证据"),
        ("协议适配层", "把 HTTP、MQTT、Matter、Modbus 等转换成统一调用", "Adapter、连接会话、错误码"),
        ("能力模型层", "描述开关、温度、亮度、锁、视频、能源等属性和动作", "能力声明、参数约束、事件类型"),
        ("控制与状态层", "命令幂等、重试、回读、期望状态与实际状态对账", "命令结果、状态快照、事件流"),
        ("应用与智能层", "前端、自动化、AI、日志、审计和第三方接口", "统一设备卡片、计划和告警"),
    ], [1900, 4700, 2760])
    add_code(doc, "设备 -> 发现证据 -> Adapter -> Capability -> Command -> 回读状态\n                         \\-> Event / Audit / Context / UI")
    add_body(doc, "前端只依赖 Capability 和状态快照；自动化只依赖事件和动作；协议差异全部封装在 Adapter 内。这是实现可扩展性的关键。")

    add_heading(doc, "3. 设备发现：不是固定动画，而是可验证的多策略探测", 1)
    add_heading(doc, "3.1 局域网发现策略", 2)
    add_table(doc, ["发现方式", "适合设备", "验证动作"], [
        ("mDNS / DNS-SD", "局域网服务、打印机、摄像头、开发板", "读取服务类型、TXT 能力、设备 ID，再请求健康接口"),
        ("SSDP / UPnP", "家电、媒体设备、路由器", "解析 LOCATION，校验设备描述和序列号"),
        ("Matter / Thread", "新一代智能家居设备", "配对码、证书链、Fabric 归属和集群能力读取"),
        ("厂商局域网目录", "没有标准广播的厂商设备", "通过主控或用户提供的目录地址读取真实在线设备"),
        ("定向地址探测", "用户已知 IP、域名或端口", "只探测允许的地址和端口，读取协议握手和状态"),
    ], [2100, 3000, 4260])
    add_heading(doc, "3.2 非局域网和物理总线发现", 2)
    add_bullet(doc, "BLE：扫描广播 UUID，连接后读取 GATT 服务、特征值、加密状态和设备序列号。")
    add_bullet(doc, "Zigbee / Thread / Z-Wave：通过协调器获取节点、端点、Cluster 或 Command Class，再映射能力。")
    add_bullet(doc, "RS-485 / Modbus：先识别总线配置，再按允许的站号范围读取设备标识和寄存器，不做无限枚举。")
    add_bullet(doc, "串口 / CAN / GPIO：由受控边缘桥接器提供设备描述和读回接口，主控不直接猜帧格式。")
    add_heading(doc, "3.3 发现结果的真实性证据", 2)
    add_table(doc, ["证据", "用途", "缺失时"], [
        ("在线握手成功", "证明地址和协议入口可达", "标记未发现或离线"),
        ("设备唯一标识", "防止重复注册和冒充", "只能保留候选，不可正式注册"),
        ("能力描述或读回状态", "证明动作确实由设备支持", "能力为未知，不允许自动执行"),
        ("认证 / 配对结果", "证明控制权限属于主控", "只读或拒绝控制"),
    ], [2500, 4500, 2360])

    add_heading(doc, "4. 统一能力模型：让不同设备变成同一种可理解对象", 1)
    add_body(doc, "不要用设备名称猜功能。每个设备必须发布一份能力描述，包含属性、动作、事件、参数范围、单位、读写权限和回读方式。一个空调可以同时有 power、mode、temperature、fan_speed；一个门锁可以有 locked、open、alarm 和 unlock 动作。")
    add_code(doc, '''{
  "device_id": "vendor-serial-001",
  "identity": {"vendor": "示例厂商", "model": "X1", "serial": "..."},
  "transport": {"protocol": "mqtt", "endpoint": "内部保存", "security": "mTLS"},
  "properties": [
    {"name": "temperature", "type": "number", "unit": "°C", "read": true},
    {"name": "power", "type": "boolean", "read": true, "write": true}
  ],
  "actions": [
    {"name": "set_power", "params": {"on": "boolean"}, "idempotent": true}
  ],
  "events": ["state_changed", "fault"]
}''')
    add_table(doc, ["能力字段", "必须说明"], [
        ("property", "当前值、数据类型、单位、精度、读写权限、更新时间"),
        ("action", "动作名、必选参数、范围、幂等性、超时、回读要求"),
        ("event", "事件名称、来源、时间戳、去重键、严重级别"),
        ("fault", "错误码、可恢复性、重试建议、用户可见文案"),
    ], [2200, 7160])

    add_heading(doc, "5. 协议适配器：一次实现，所有页面和智能能力复用", 1)
    add_heading(doc, "5.1 适配器生命周期", 2)
    for item in [
        "声明 adapter_id、支持的协议、设备类型过滤器和版本范围。",
        "执行 discover / identify，返回真实设备身份和发现证据。",
        "执行 describe，返回能力模型；未知能力不进入可执行列表。",
        "创建安全连接会话，保存令牌引用而不是把密钥写进设备记录。",
        "实现 read_state、invoke、subscribe_events 和 close。",
        "报告健康状态、延迟、重试次数和最后一次成功回读时间。",
    ]:
        add_number(doc, item)
    add_heading(doc, "5.2 适配器最小接口", 2)
    add_code(doc, '''class DeviceAdapter:
    adapter_id = "vendor_protocol_v1"

    def discover(self, scope): ...
    def identify(self, candidate): ...
    def describe(self, identity): ...
    def connect(self, identity, secret_ref): ...
    def read_state(self, session, device): ...
    def invoke(self, session, device, action, params, request_id): ...
    def subscribe_events(self, session, device, callback): ...
    def close(self, session): ...''')
    add_body(doc, "适配器必须是可测试、可超时、可取消的插件。它不能直接修改数据库、UI 或 AI 对话；所有结果通过统一事件和状态接口返回。")

    add_heading(doc, "6. 主流协议适配矩阵", 1)
    add_table(doc, ["协议 / 总线", "连接模型", "典型能力", "安全与适配重点"], [
        ("HTTPS / REST", "请求-响应", "灯、空调、摄像头、网关", "TLS、证书校验、签名、超时、幂等键"),
        ("MQTT", "发布-订阅", "传感器、开关、状态事件", "TLS、客户端证书、主题 ACL、消息去重"),
        ("WebSocket", "长连接双向事件", "实时状态、报警、面板", "握手认证、心跳、断线重连、帧完整性"),
        ("CoAP / DTLS", "低功耗资源模型", "电池传感器、照明", "资源白名单、重放保护、确认消息"),
        ("Matter", "标准集群与 Fabric", "照明、插座、门锁、环境", "配对码、证书、Fabric 权限、Cluster 映射"),
        ("Modbus TCP / RTU", "寄存器读写", "电表、空调、工业设备", "站号白名单、寄存器 schema、CRC、写入保护"),
        ("BLE GATT", "短距服务特征", "穿戴、锁、传感器", "配对、特征权限、连接超时、重连"),
        ("Zigbee / Thread / Z-Wave", "协调器节点", "低功耗家居设备", "网络密钥、节点权限、端点和 Cluster 版本"),
        ("RS-485 / CAN / GPIO", "物理总线或边缘桥接", "继电器、工业传感器", "总线隔离、帧校验、速率、驱动白名单"),
    ], [1800, 2000, 2600, 2960])
    add_body(doc, "协议越多，越需要统一的能力模型和错误语义。增加协议不等于把所有端口暴露给前端；协议网关应按设备、动作和网络范围建立最小权限。")

    add_heading(doc, "7. 安全设计：接入任意设备，但不能放大风险", 1)
    add_table(doc, ["安全层", "实施要求"], [
        ("边界", "仅允许用户选择的网络范围、设备类型和端口；拒绝公网任意探测和开放式端口扫描"),
        ("身份", "设备唯一 ID、证书 / 配对码 / 令牌引用；防止重复注册和设备冒充"),
        ("传输", "优先 TLS / DTLS / 安全总线；旧 HTTP 只限隔离内网并配合签名、时间戳和重放保护"),
        ("授权", "按设备、动作、用户、自动化来源建立 ACL；开锁、写寄存器、固件升级单独授权"),
        ("密钥", "密钥进入安全存储或系统密钥环，数据库和前端只保存 secret_ref，不记录明文"),
        ("执行", "参数 schema 校验、速率限制、幂等键、超时、熔断和审计日志"),
        ("供应链", "适配器签名、版本锁定、沙箱权限、依赖扫描和回滚包"),
    ], [1800, 7560])
    add_body(doc, "安全边界必须优先于“自动接入”。如果设备的协议无法认证或无法验证状态，应当允许只读观察或明确拒绝，而不是为了演示而伪造在线和成功。")

    add_heading(doc, "8. 状态同步：解决“按钮变了但设备没变”", 1)
    add_heading(doc, "8.1 期望状态与实际上报状态", 2)
    add_body(doc, "控制请求只代表“主控发出了意图”，不能直接当成设备已经完成。系统应保存 desired_state、reported_state、request_id、device_timestamp 和 confirmed_at。只有回读或设备事件确认后，UI 才显示完成。")
    add_heading(doc, "8.2 可靠执行流程", 2)
    for item in [
        "生成全局唯一 request_id 和幂等键，检查动作参数。",
        "写入命令日志，状态标记为 pending。",
        "通过适配器发送动作，带上超时、认证和重试策略。",
        "读取设备回执或主动回读状态，校验值、时间戳和设备 ID。",
        "一致则标记 confirmed；不一致则标记 mismatch / failed，并显示具体原因。",
        "离线时进入队列或明确拒绝，恢复上线后按策略重新对账，不重复执行非幂等动作。",
    ]:
        add_number(doc, item)
    add_table(doc, ["状态", "含义", "前端表现"], [
        ("pending", "命令已发出，等待设备确认", "显示处理中，不假装开关已完成"),
        ("confirmed", "回读与期望一致", "显示目标状态和完成时间"),
        ("mismatch", "设备返回与目标不一致", "显示实际值、原因和重试入口"),
        ("offline", "设备或链路不可达", "显示离线，不允许危险动作"),
    ], [1800, 3900, 3660])

    add_heading(doc, "9. 自动化、AI 与自定义设备的边界", 1)
    add_body(doc, "智能层只能从能力模型中选择动作，不能绕过适配器直接拼接厂商命令。AI 负责理解意图、排序候选、解释原因和提出计划；执行器负责权限、参数校验、并发、回读和安全策略。")
    add_bullet(doc, "只读信息可以自动汇总；开锁、门禁、燃气、电源总闸、固件升级等高风险动作必须走独立授权策略。")
    add_bullet(doc, "自动化规则以事件、阈值、时间窗和设备能力为条件，不以设备昵称或模糊字符串直接执行。")
    add_bullet(doc, "每次自动动作都写入原因、输入状态、命令、回读结果、失败原因和策略版本，便于评价和回滚。")
    add_bullet(doc, "新增设备未声明的能力不会被 AI 猜测；用户可以先使用只读模式，再逐项开放写入能力。")

    add_heading(doc, "10. 插件清单与设备注册示例", 1)
    add_code(doc, '''{
  "adapter_id": "acme.modbus.v2",
  "version": "2.1.0",
  "match": {"protocol": "modbus_rtu", "vendor": "ACME", "model": ["X1", "X2"]},
  "transport": {"baudrate": 9600, "parity": "N", "stop_bits": 1},
  "capabilities": [
    {"name": "power", "type": "boolean", "read_register": 1, "write_register": 1},
    {"name": "temperature", "type": "number", "unit": "°C", "scale": 0.1, "read_register": 10}
  ],
  "permissions": ["read_state", "write_power"],
  "health": {"heartbeat_seconds": 30, "read_timeout_ms": 800}
}''')
    add_body(doc, "清单由适配器作者提供，主控在安装前验证签名、版本、权限和能力 schema。设备注册时只写入身份、能力、连接引用和策略，不把厂商密钥复制到手机端。")

    add_heading(doc, "11. 端到端实施流程", 1)
    for item in [
        "为目标设备选择连接方式：标准协议、厂商 API、边缘桥接或物理总线。",
        "编写设备描述文件，列出身份字段、属性、动作、事件、错误码和安全要求。",
        "实现 Adapter，并用离线模拟器、真实设备和异常链路分别测试。",
        "将 Adapter 安装到协议网关，限制它能访问的地址、端口、文件和系统调用。",
        "执行发现与身份认证，保留发现证据；失败时返回可解释原因。",
        "读取能力模型，生成统一设备卡片和 API；未知能力保持禁用。",
        "进行一次只读状态同步，再开放低风险写入动作。",
        "按 desired / reported 状态模型接入自动化、AI、日志和移动端。",
        "完成断网、重启、重复命令、设备替换、权限撤销和回滚验收。",
    ]:
        add_number(doc, item)

    add_heading(doc, "12. 测试与验收矩阵", 1)
    add_table(doc, ["测试项", "验收标准"], [
        ("离线发现", "设备断电或不可达时不出现在已接入列表，无假成功"),
        ("身份唯一性", "同一设备重复扫描不会重复注册，替换设备不会继承旧密钥"),
        ("能力准确性", "页面只显示适配器声明并通过读回验证的动作"),
        ("控制闭环", "命令、回执、回读、超时和错误均有 request_id 关联"),
        ("协议安全", "认证失败、证书错误、重放、越权和公网地址均被拒绝"),
        ("断网恢复", "恢复后自动重新对账，非幂等命令不会重复执行"),
        ("并发控制", "多个来源同时操作时有版本或锁策略，最终状态可解释"),
        ("审计追踪", "可按设备、用户、适配器、时间和 request_id 查询完整记录"),
        ("回滚", "适配器或设备描述升级失败可恢复旧版本，不破坏现有设备"),
    ], [2400, 6960])

    add_heading(doc, "13. 运营、升级和扩展建议", 1)
    add_bullet(doc, "适配器版本与设备描述版本分离管理，支持灰度发布和按设备回滚。")
    add_bullet(doc, "为每种协议维护连接池、心跳、限流、断线重连和指标：成功率、延迟、离线时长、回读一致率。")
    add_bullet(doc, "建立厂商适配器认证流程：协议文档、抓包样例、异常码、最小权限和现场验收记录缺一不可。")
    add_bullet(doc, "把设备能力目录作为长期资产，前端、自动化和 AI 都从目录读取，不复制一份私有别名表。")
    add_bullet(doc, "当设备协议发生变化时，优先升级 Adapter，不改业务层和页面；这是控制平台可持续演进的核心。")

    add_heading(doc, "14. 结论", 1)
    add_body(doc, "一个真正可扩展的自定义设备平台，不是把几个设备类型写在页面上，而是建立“真实发现 + 身份认证 + 能力描述 + 协议适配 + 状态闭环 + 安全审计”的完整工程链路。这样接入灯、空调、门锁、摄像头、传感器、工业仪表或自研开发板时，新增的是设备描述和适配器，核心控制平面、前端组件、日志和智能能力都可以复用。")
    add_body(doc, "最终判断标准只有一个：设备在线时能被证明、能力能被读取、命令能被执行、结果能被回读；设备不在线或不可信时，系统必须诚实地报告失败。")
    return doc


def main():
    for output in OUTPUTS:
        output.parent.mkdir(parents=True, exist_ok=True)
        make_doc().save(output)
        print(output)


if __name__ == "__main__":
    main()
